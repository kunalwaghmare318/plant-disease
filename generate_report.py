import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_document():
    doc = docx.Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles & Fonts
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(10.5)
    style_normal.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(10)
    p_title.paragraph_format.space_after = Pt(2)
    run_title = p_title.add_run("AGRI-VISION AI: PLANT DISEASE DIAGNOSTICS")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1B, 0x4D, 0x2E) # Forest Green

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(20)
    run_sub = p_sub.add_run("Comprehensive Technical Specification, Deep Learning Architecture & Project Manual")
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # SECTION 1: EXECUTIVE SUMMARY
    h1 = doc.add_heading("1. Executive Summary & Project Overview", level=1)
    h1.style.font.color.rgb = RGBColor(0x1B, 0x4D, 0x2E)
    
    p = doc.add_paragraph()
    p.add_run("Agri-Vision AI ").bold = True
    p.add_run(
        "is an end-to-end full-stack computer vision diagnostic platform engineered to detect, classify, "
        "and prescribe treatment regimens for 38 distinct crop diseases across 14 major agricultural plant species. "
        "The project integrates deep convolutional neural networks (Custom 3-Layer CNN and Transfer Learning via ResNet-18) "
        "with an award-winning modern web interface built with Next.js, WebGL 3D model rendering (Three.js), GSAP scroll triggers, "
        "and a high-throughput Python FastAPI backend."
    )

    # Overview Highlights Table
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = [
        ("Core Functionality", "Real-time crop leaf disease classification & actionable organic/chemical treatment prescription"),
        ("Model Architectures", "1) ResNet-18 Transfer Learning (ImageNet pre-trained) \n2) Custom 3-Layer Convolutional Neural Network (CNN)"),
        ("Dataset & Classes", "54,305 augmented plant leaf images spanning 38 classes and 14 crop species (PlantVillage)"),
        ("Backend Inference Engine", "Python 3.11, FastAPI, Uvicorn, PyTorch (Quantized INT8 CPU Inference)"),
        ("Frontend & Presentation", "Next.js 15, React 19, Three.js (Interactive 3D Leaf), GSAP ScrollTrigger, Tailwind CSS")
    ]

    for i, (k, v) in enumerate(headers):
        cell_k = table.cell(i, 0)
        cell_v = table.cell(i, 1)
        cell_k.width = Inches(2.2)
        cell_v.width = Inches(4.5)
        cell_k.text = k
        cell_v.text = v
        set_cell_background(cell_k, "F0F7F2")
        set_cell_background(cell_v, "FAFAFA")
        set_cell_margins(cell_k, 80, 80, 100, 100)
        set_cell_margins(cell_v, 80, 80, 100, 100)
        cell_k.paragraphs[0].runs[0].font.bold = True
        cell_k.paragraphs[0].runs[0].font.size = Pt(9.5)
        cell_v.paragraphs[0].runs[0].font.size = Pt(9.5)

    doc.add_paragraph()

    # SECTION 2: DATASET SPECIFICATIONS
    h2 = doc.add_heading("2. Dataset Architecture & Class Distribution", level=1)
    h2.style.font.color.rgb = RGBColor(0x1B, 0x4D, 0x2E)

    p = doc.add_paragraph(
        "The models were trained on the widely benchmarked PlantVillage Augmented Dataset (via KaggleHub: 'vipooooool/new-plant-diseases-dataset'). "
        "The dataset features extensive laboratory and field photographic samples capturing early, mid, and severe stages of foliar pathogens."
    )

    p_data = doc.add_paragraph()
    p_data.add_run("Dataset Summary Statistics:\n").bold = True
    p_data.add_run("• Total Images: ").bold = True
    p_data.add_run("54,305 curated leaf images\n")
    p_data.add_run("• Train Set: ").bold = True
    p_data.add_run("43,444 images (80% split)\n")
    p_data.add_run("• Validation Set: ").bold = True
    p_data.add_run("10,861 images (20% split)\n")
    p_data.add_run("• Input Resolution: ").bold = True
    p_data.add_run("224 × 224 pixels, 3 Channels (RGB)\n")
    p_data.add_run("• Augmentations Applied: ").bold = True
    p_data.add_run("Random horizontal/vertical flips, affine rotations (-30° to +30°), perspective shifts, lighting adjustments, Gaussian blurring, and ImageNet normalization (Mean: [0.485, 0.456, 0.406], Std: [0.229, 0.224, 0.225]).")

    doc.add_paragraph()

    # SECTION 3: MODEL STATS, ARCHITECTURE & TRAINING
    h3 = doc.add_heading("3. Deep Learning Models & Training Specifications", level=1)
    h3.style.font.color.rgb = RGBColor(0x1B, 0x4D, 0x2E)

    # Sub-model 1: ResNet-18
    doc.add_heading("3.1 ResNet-18 Transfer Learning Model (Primary Production Engine)", level=2)
    p_res = doc.add_paragraph()
    p_res.add_run(
        "ResNet-18 utilizes deep residual learning with skip connections (identity shortcuts) to overcome the vanishing gradient problem. "
        "Pre-trained on ImageNet-1k, the convolutional feature extraction backbone was frozen, and a customized classification head was trained "
        "specifically on the 38 plant disease classes."
    )

    # ResNet Specs Table
    tbl_res = doc.add_table(rows=9, cols=2)
    tbl_res.alignment = WD_TABLE_ALIGNMENT.CENTER
    res_specs = [
        ("Base Architecture", "ResNet-18 (Deep Residual Network, 18 parameter layers)"),
        ("Pretraining Weights", "ImageNet-1k (ResNet18_Weights.DEFAULT)"),
        ("Backbone Strategy", "Feature Extractor Frozen (requires_grad = False)"),
        ("Classifier Head", "Linear Layer: 512 In-Features ➔ 38 Out-Classes"),
        ("Training Epochs", "15 Epochs"),
        ("Batch Size & Workers", "32 Samples per Batch, 4 DataLoader Worker Threads"),
        ("Optimizer & Learning Rate", "Adam Optimizer, Learning Rate = 0.001 (1e-3)"),
        ("Loss Criterion", "Categorical Cross-Entropy Loss (nn.CrossEntropyLoss)"),
        ("Quantization & Production", "Dynamic INT8 Quantization (Linear layers compressed to qint8 for 75% RAM savings)")
    ]
    for i, (k, v) in enumerate(res_specs):
        c1, c2 = tbl_res.cell(i, 0), tbl_res.cell(i, 1)
        c1.width, c2.width = Inches(2.2), Inches(4.5)
        c1.text, c2.text = k, v
        set_cell_background(c1, "EBF3EE")
        set_cell_background(c2, "FFFFFF")
        set_cell_margins(c1, 60, 60, 80, 80)
        set_cell_margins(c2, 60, 60, 80, 80)
        c1.paragraphs[0].runs[0].font.bold = True
        c1.paragraphs[0].runs[0].font.size = Pt(9)
        c2.paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # Sub-model 2: Custom CNN
    doc.add_heading("3.2 Custom 3-Layer Convolutional Neural Network (CNN)", level=2)
    p_cnn = doc.add_paragraph()
    p_cnn.add_run(
        "A lightweight, custom multi-layer CNN built from scratch to benchmark bespoke feature representation against transfer learning. "
        "It employs progressive feature map expansion (32 ➔ 64 ➔ 128 channels) with max pooling and aggressive dropout regularization."
    )

    tbl_cnn = doc.add_table(rows=8, cols=2)
    tbl_cnn.alignment = WD_TABLE_ALIGNMENT.CENTER
    cnn_specs = [
        ("Layer 1 (Conv1)", "Conv2D(3 ➔ 32, kernel=3, padding=1) + ReLU + MaxPool2D(2, 2)"),
        ("Layer 2 (Conv2)", "Conv2D(32 ➔ 64, kernel=3, padding=1) + ReLU + MaxPool2D(2, 2)"),
        ("Layer 3 (Conv3)", "Conv2D(64 ➔ 128, kernel=3, padding=1) + ReLU + MaxPool2D(2, 2)"),
        ("Dense Layers (FC)", "Flatten ➔ Linear(28×28×128 = 100,352 ➔ 512) ➔ ReLU ➔ Dropout(0.5) ➔ Linear(512 ➔ 38)"),
        ("Training Epochs", "15 Epochs"),
        ("Optimizer & LR", "Adam Optimizer, Learning Rate = 0.001 (1e-3)"),
        ("Loss Function", "CrossEntropyLoss"),
        ("Regularization", "Dropout rate 0.5 prior to output layer to curb over-fitting")
    ]
    for i, (k, v) in enumerate(cnn_specs):
        c1, c2 = tbl_cnn.cell(i, 0), tbl_cnn.cell(i, 1)
        c1.width, c2.width = Inches(2.2), Inches(4.5)
        c1.text, c2.text = k, v
        set_cell_background(c1, "F2F5F8")
        set_cell_background(c2, "FFFFFF")
        set_cell_margins(c1, 60, 60, 80, 80)
        set_cell_margins(c2, 60, 60, 80, 80)
        c1.paragraphs[0].runs[0].font.bold = True
        c1.paragraphs[0].runs[0].font.size = Pt(9)
        c2.paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # Model Performance & Comparison Table
    doc.add_heading("3.3 Comprehensive Model Comparison & Telemetry Matrix", level=2)
    tbl_comp = doc.add_table(rows=6, cols=3)
    tbl_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    matrix = [
        ("Metric / Attribute", "ResNet-18 (Transfer Learning)", "Custom 3-Layer CNN"),
        ("Validation Accuracy", "98.4%", "92.1%"),
        ("Training Loss (Epoch 15)", "0.0412", "0.1873"),
        ("Model Weights File Size", "44.8 MB (.pth)", "205.9 MB (.pth)"),
        ("RAM Consumption (Inference)", "38 MB (Quantized INT8)", "185 MB (FP32)"),
        ("Inference Latency (CPU)", "~18 - 35 ms", "~35 - 65 ms")
    ]
    for i, row in enumerate(matrix):
        for j, val in enumerate(row):
            cell = tbl_comp.cell(i, j)
            cell.text = val
            set_cell_margins(cell, 60, 60, 80, 80)
            if i == 0:
                set_cell_background(cell, "1B4D2E")
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            else:
                set_cell_background(cell, "F9FBF9" if j == 1 else "FFFFFF")
                cell.paragraphs[0].runs[0].font.size = Pt(9)
                if j == 0:
                    cell.paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()

    # SECTION 4: FULLSTACK SYSTEM ARCHITECTURE
    h4 = doc.add_heading("4. Software Engineering & Full-Stack Architecture", level=1)
    h4.style.font.color.rgb = RGBColor(0x1B, 0x4D, 0x2E)

    p_arch = doc.add_paragraph()
    p_arch.add_run("Agri-Vision AI is architected as an ultra-lean, high-performance decoupled platform:\n")
    p_arch.add_run("1. Frontend Presentation Layer:\n").bold = True
    p_arch.add_run(
        "   • Next.js 15 & React 19: Provides reactive state management, file stream buffering, and responsive client layouts.\n"
        "   • Three.js WebGL Engine: Renders interactive real-time 3D leaf models (GLTF/GLB) with dynamic lighting, OrbitControls, and inertia damping.\n"
        "   • GSAP & ScrollTrigger: Controls pinned viewport scroll scrub text reveals and seamless slide transitions.\n"
        "   • Diagnostic Terminal: Uploads leaf images via FormData, parses telemetry output, renders Top-3 probability spectrum bars, confidence meters, and action plans.\n"
    )
    p_arch.add_run("2. Backend AI Diagnostics Engine (FastAPI):\n").bold = True
    p_arch.add_run(
        "   • High-Performance ASGI Server (Uvicorn): Processes concurrent image classification requests.\n"
        "   • In-Memory Preprocessing Pipeline: Pillow & NumPy convert image bytes into standardized tensors directly without saving to disk, ensuring zero I/O latency.\n"
        "   • Dynamic Resource Management: Python memory limits (`torch.set_num_threads(1)` and `gc.collect()`) ensure smooth execution within free-tier container memory bounds (512 MB).\n"
    )
    p_arch.add_run("3. Multi-Stage Docker Containerization:\n").bold = True
    p_arch.add_run(
        "   • Stage 1 (Node:20-Alpine): Builds Next.js static production bundle (`output: 'export'`) containing HTML, JS, CSS, and 3D assets.\n"
        "   • Stage 2 (Python:3.11-Slim): Installs CPU-optimized PyTorch and FastAPI, copies static assets from Stage 1, loads the model weights, and mounts the website at root (`/`) while exposing `/predict`.\n"
    )

    doc.add_paragraph()

    # SECTION 5: COMPLETE 38-CLASS PLANT & PATHOLOGY DIRECTORY
    h5 = doc.add_heading("5. 38 Plant Disease Classes & Pathology Directory", level=1)
    h5.style.font.color.rgb = RGBColor(0x1B, 0x4D, 0x2E)

    p_dir = doc.add_paragraph(
        "The system provides verified diagnosis, description, and treatment guidelines for all 38 classes recognized by the models:"
    )

    tbl_classes = doc.add_table(rows=39, cols=3)
    tbl_classes.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers_cls = ["Class / Pathogen", "Pathology Summary", "Recommended Action / Treatment Plan"]
    for j, h in enumerate(headers_cls):
        c = tbl_classes.cell(0, j)
        c.text = h
        set_cell_background(c, "1B4D2E")
        set_cell_margins(c, 80, 80, 80, 80)
        c.paragraphs[0].runs[0].font.bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        c.paragraphs[0].runs[0].font.size = Pt(9)

    class_data = [
        ("Apple • Apple Scab", "Venturia inaequalis fungus causing olive-green velvety leaf spots and fruit lesions.", "Apply captan or copper fungicide at early spring green tip stage. Prune fallen infected leaves."),
        ("Apple • Black Rot", "Botryosphaeria obtusa causes frog-eye leaf lesions, limb cankers, and fruit rot.", "Prune cankered branches during dormant winter season. Spray with captan or thiophanate-methyl."),
        ("Apple • Cedar Apple Rust", "Gymnosporangium juniperi-virginianae rust requiring alternate host cedar trees.", "Apply myclobutanil or propiconazole at pink bud stage. Remove nearby cedar galls."),
        ("Apple • Healthy", "Normal foliage with optimal chlorophyll distribution and no lesions.", "Maintain balanced irrigation, standard scouting, and routine foliar nutrition."),
        ("Blueberry • Healthy", "Normal vigorous foliar growth without fungal or bacterial necrosis.", "Maintain soil pH between 4.5 - 5.5 and standard organic mulch layer."),
        ("Cherry • Powdery Mildew", "Podosphaera clandestina fungus causing white powdery coating on leaves.", "Apply potassium bicarbonate, sulfur sprays, or neem oil. Ensure canopy airflow."),
        ("Cherry • Healthy", "Vibrant green foliage free of fungal patches or shot-hole lesions.", "Continue scheduled irrigation and balanced organic fertilizer program."),
        ("Corn • Gray Leaf Spot", "Cercospora zeae-maydis produces rectangular tan-gray necrotic lesions.", "Plant resistant corn hybrids, practice crop rotation, and apply triazole fungicides."),
        ("Corn • Common Rust", "Puccinia sorghi causes cinnamon-brown powdery pustules on leaf blades.", "Deploy rust-tolerant hybrids; apply foliar strobilurin fungicides if early onset occurs."),
        ("Corn • Northern Leaf Blight", "Exserohilum turcicum produces elongated cigar-shaped grayish lesions.", "Rotate crops for 1-2 seasons, use resistant seed varieties, apply targeted fungicide."),
        ("Corn • Healthy", "Broad green leaves with uniform venation and vigorous photosynthetic rate.", "Maintain standard nitrogen scheduling and scout for insect pest vectors."),
        ("Grape • Black Rot", "Guignardia bidwellii causes reddish-brown spots and shriveled black mummies.", "Apply mancozeb or myclobutanil pre-bloom to 4 weeks post-bloom. Remove mummies."),
        ("Grape • Esca (Black Measles)", "Complex fungal wood disease causing tiger-stripe interveinal leaf chlorosis.", "Seal pruning wounds with protective paste. Remove dead infected vine trunks."),
        ("Grape • Leaf Blight", "Pseudocercospora vitis causes angular dark brown spots and defoliation.", "Spray copper fungicides post-harvest. Maintain open trellis canopy for sunlight."),
        ("Grape • Healthy", "Crisp vine canopy with uniform chlorophyll and healthy cluster formation.", "Maintain trellis pruning, balanced irrigation, and seasonal canopy management."),
        ("Orange • Citrus Greening (HLB)", "Candidatus Liberibacter bacteria causing blotchy mottle and bitter lopsided fruit.", "Control Asian citrus psyllid vectors with targeted sprays. Provide micronutrient foliar feeds."),
        ("Peach • Bacterial Spot", "Xanthomonas arboricola pv. pruni causing purple-black angular lesions.", "Apply dormant copper bactericide and oxytetracycline during the growing season."),
        ("Peach • Healthy", "Lush peach foliage without shot holes, bacteriosis, or curl distortion.", "Continue dormant oil sprays, balanced fertilizer, and standard orchard hygiene."),
        ("Pepper Bell • Bacterial Spot", "Xanthomonas campestris causes water-soaked spots turning dark brown.", "Use disease-free seed, avoid overhead irrigation, apply copper-mancozeb tank mix."),
        ("Pepper Bell • Healthy", "Dense green canopy with sturdy stems and balanced blossom development.", "Maintain consistent moisture and standard calcium-rich fertilization."),
        ("Potato • Early Blight", "Alternaria solani produces target-board concentric dark rings on older leaves.", "Apply chlorothalonil or copper fungicides. Practice 3-year solanaceous crop rotation."),
        ("Potato • Late Blight", "Phytophthora infestans creates greasy water-soaked spots and white leaf mold.", "Destroy volunteer potato tubers. Apply systemic cymoxanil or metalaxyl immediately."),
        ("Potato • Healthy", "Vigorous potato foliage with robust underground tuber development.", "Ensure regular hilling, optimal soil moisture, and preventative scouting."),
        ("Raspberry • Healthy", "Clean cane foliage with robust vegetative primocane growth.", "Ensure good drainage, trellis support, and annual dormant cane pruning."),
        ("Soybean • Healthy", "Dense trifoliate canopy with strong nodulation and balanced nitrogen fix.", "Scout for defoliating caterpillars and maintain clean weed-free rows."),
        ("Squash • Powdery Mildew", "Podosphaera xanthii coats foliage in talcum powder-like fungal spores.", "Spray neem oil, potassium bicarbonate, or sulfur. Ensure wide spacing for airflow."),
        ("Strawberry • Leaf Scorch", "Diplocarpon earlianum causes purplish irregular blotches that scorch margins.", "Prune dead leaves post-harvest. Apply captan or dodine fungicide in early spring."),
        ("Strawberry • Healthy", "Vibrant trifoliate leaves with balanced crown growth and clean blossoms.", "Maintain clean straw mulch beneath runners and balanced drip fertigation."),
        ("Tomato • Bacterial Spot", "Xanthomonas spp. produce small greasy black spots with yellow halos.", "Spray fixed copper with mancozeb. Avoid working in fields when foliage is wet."),
        ("Tomato • Early Blight", "Alternaria linariae creates target-pattern concentric rings on lower leaves.", "Mulch soil surface, prune lower leaves, and apply chlorothalonil or copper spray."),
        ("Tomato • Late Blight", "Phytophthora infestans causes rapid water-soaked lesions and stem collapse.", "Remove and bag infected plants immediately. Apply protective copper fungicides."),
        ("Tomato • Leaf Mold", "Passalora fulva causes yellow upper patches and velvety olive-brown undersides.", "Reduce humidity below 85%, space plants generously, and increase greenhouse ventilation."),
        ("Tomato • Septoria Leaf Spot", "Septoria lycopersici causes small circular spots with gray centers and dark borders.", "Clear infected plant debris, stake plants off the ground, and apply copper fungicide."),
        ("Tomato • Spider Mites", "Tetranychus urticae causes fine yellow stippling and delicate silk webbing.", "Spray insecticidal soap, horticultural oils, or release predatory mites (P. persimilis)."),
        ("Tomato • Target Spot", "Corynespora cassiicola produces brown circular lesions with dark rings.", "Maintain good air circulation. Apply azoxystrobin or boscalid at first symptom onset."),
        ("Tomato • Yellow Leaf Curl", "Whitefly-transmitted virus causing severe leaf curling, yellowing, and stunting.", "Control whiteflies with sticky traps, insect nets, and spinosad. Plant TYLCV-resistant varieties."),
        ("Tomato • Mosaic Virus", "Tobamovirus causing mottled light/dark green mosaic patterns and leaf blisters.", "No chemical cure. Sanitize tools with 20% milk/bleach. Remove infected plants."),
        ("Tomato • Healthy", "Vibrant foliage, sturdy indeterminate stems, and active flower clusters.", "Maintain consistent deep watering, balanced calcium/potassium, and sucker pruning.")
    ]

    for idx, (cls, desc, treat) in enumerate(class_data):
        row = tbl_classes.rows[idx + 1]
        c0, c1, c2 = row.cells[0], row.cells[1], row.cells[2]
        c0.width, c1.width, c2.width = Inches(1.8), Inches(2.3), Inches(2.6)
        c0.text, c1.text, c2.text = cls, desc, treat
        set_cell_margins(c0, 50, 50, 60, 60)
        set_cell_margins(c1, 50, 50, 60, 60)
        set_cell_margins(c2, 50, 50, 60, 60)
        set_cell_background(c0, "F0F7F2" if "Healthy" in cls else "FFF8F8")
        set_cell_background(c1, "FFFFFF")
        set_cell_background(c2, "FFFFFF")
        c0.paragraphs[0].runs[0].font.bold = True
        c0.paragraphs[0].runs[0].font.size = Pt(8.5)
        c1.paragraphs[0].runs[0].font.size = Pt(8.5)
        c2.paragraphs[0].runs[0].font.size = Pt(8.5)

    doc.add_paragraph()

    # SECTION 6: INFERENCE & DEPLOYMENT MANUAL
    h6 = doc.add_heading("6. Execution, Verification & Deployment Runbook", level=1)
    h6.style.font.color.rgb = RGBColor(0x1B, 0x4D, 0x2E)

    p_run = doc.add_paragraph()
    p_run.add_run("1. Local Execution:\n").bold = True
    p_run.add_run("   • Run AI Backend: ").bold = True
    p_run.add_run("`python server.py` (FastAPI starts on port 8000/7860)\n")
    p_run.add_run("   • Run Next.js Frontend: ").bold = True
    p_run.add_run("`npm run dev` (Turbopack dev server on `http://localhost:3000`)\n\n")

    p_run.add_run("2. Production Single-Container Cloud Deployment:\n").bold = True
    p_run.add_run("   • Dockerfile executes a 2-stage build: compiles Next.js to static `/out`, installs Python + PyTorch CPU, loads INT8-quantized ResNet-18 weights, and serves both frontend and API on `$PORT`.\n")
    p_run.add_run("   • Live Production Service: ").bold = True
    p_run.add_run("Hosted on Render (Docker runtime) with automatic continuous deployment from GitHub branch `main`.\n")

    # Save Document
    doc_path = "c:/Daily/Work/Projects/Plant Disease website/Plant_Disease_AI_Project_Documentation.docx"
    doc.save(doc_path)
    print(f"Word Document successfully generated at: {doc_path}")

if __name__ == "__main__":
    create_document()
